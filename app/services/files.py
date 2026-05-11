"""Document loading utilities."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
import csv
import io
import re
import subprocess
from typing import Any

from langchain_core.documents import Document

from app.core.settings import settings

# Common plain-text formats that can be loaded with UTF-8 fallback.
TEXT_FILE_EXTENSIONS: tuple[str, ...] = (
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".tsv",
    ".json",
    ".jsonl",
    ".yaml",
    ".yml",
    ".xml",
    ".ini",
    ".cfg",
    ".conf",
    ".toml",
    ".log",
    ".rst",
    ".rtf",
    ".sql",
    ".py",
    ".js",
    ".ts",
    ".jsx",
    ".tsx",
    ".java",
    ".c",
    ".cpp",
    ".h",
    ".hpp",
    ".go",
    ".rs",
    ".sh",
    ".bat",
    ".ps1",
)

MARKDOWN_EXTENSIONS: tuple[str, ...] = (".md", ".markdown")
SPREADSHEET_EXTENSIONS: tuple[str, ...] = (".xlsx", ".xlsm", ".xls")
RICH_DOC_EXTENSIONS: tuple[str, ...] = (".pdf", ".doc", ".docx", *SPREADSHEET_EXTENSIONS)
TABLE_VIEW_MAX_ROWS = 600
TABLE_VIEW_MAX_COLUMNS = 50


def iter_source_files() -> list[Path]:
    """Scan data/docs and data/uploads for supported files."""
    files: list[Path] = []
    for base_dir in (settings.docs_dir, settings.uploads_dir):
        if not base_dir.exists():
            continue
        for path in sorted(base_dir.rglob("*")):
            if path.is_file() and path.suffix.lower() in settings.supported_extensions:
                files.append(path)
    return files


def file_info(path: Path) -> dict[str, str | int]:
    """Build metadata for frontend file list."""
    stat = path.stat()
    return {
        "path": str(path.relative_to(settings.root_dir)).replace("\\", "/"),
        "size_bytes": stat.st_size,
        "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds"),
        "extension": path.suffix.lower(),
    }


def _relative_source(path: Path) -> str:
    return str(path.relative_to(settings.root_dir)).replace("\\", "/")


def _safe_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return f"{value:g}"
    return str(value).strip()


def _escape_table_cell(value: str) -> str:
    return value.replace("\r", " ").replace("\n", " ").replace("|", r"\|").strip()


def _normalize_tabular_rows(rows: list[list[str]]) -> dict[str, Any]:
    width = max((len(row) for row in rows), default=0)
    if width <= 0:
        return {
            "headers": [],
            "body": [],
            "header_row_number": 0,
            "data_start_row": 1,
            "column_count": 0,
            "row_count": 0,
        }

    normalized = [row + [""] * (width - len(row)) for row in rows]
    first_row = normalized[0]
    use_first_as_header = any(cell.strip() for cell in first_row)
    if use_first_as_header:
        headers = [cell or f"Column {index + 1}" for index, cell in enumerate(first_row)]
        body = normalized[1:]
        header_row_number = 1
        data_start_row = 2
    else:
        headers = [f"Column {index + 1}" for index in range(width)]
        body = normalized
        header_row_number = 0
        data_start_row = 1

    return {
        "headers": headers,
        "body": body,
        "header_row_number": header_row_number,
        "data_start_row": data_start_row,
        "column_count": width,
        "row_count": len(rows),
    }


def _rows_to_markdown_table(rows: list[list[str]], title: str | None = None) -> str:
    lines: list[str] = []
    if title:
        lines.append(f"### {title}")
        lines.append("")

    normalized = _normalize_tabular_rows(rows)
    headers = normalized["headers"]
    body = normalized["body"]

    if not headers:
        lines.append("_No rows in this sheet._")
        return "\n".join(lines).strip()

    lines.append("| " + " | ".join(_escape_table_cell(cell) for cell in headers) + " |")
    lines.append("| " + " | ".join("---" for _ in headers) + " |")

    for row in body:
        lines.append("| " + " | ".join(_escape_table_cell(cell) for cell in row) + " |")

    return "\n".join(lines).strip()


def _rows_to_plain_records(rows: list[list[str]], sheet_name: str) -> str:
    lines: list[str] = [f"Sheet: {sheet_name}"]
    normalized = _normalize_tabular_rows(rows)
    headers = normalized["headers"]
    body = normalized["body"]

    if not headers:
        lines.append("(empty)")
        return "\n".join(lines)

    for row_index, row in enumerate(body, start=1):
        pairs = []
        for col_index, cell in enumerate(row):
            if not cell.strip():
                continue
            key = headers[col_index] if col_index < len(headers) else f"column_{col_index + 1}"
            pairs.append(f"{key}={cell}")
        if not pairs:
            continue
        lines.append(f"Row {row_index}: " + "; ".join(pairs))

    if len(lines) == 1:
        lines.append("(empty)")
    return "\n".join(lines)


def _extract_docx_markdown(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("Missing dependency: python-docx") from exc

    doc = DocxDocument(str(path))
    sections: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        if style_name.startswith("heading"):
            match = re.search(r"(\d+)", style_name)
            level = int(match.group(1)) if match else 2
            level = max(1, min(level + 1, 6))
            sections.append(f"{'#' * level} {text}")
        else:
            sections.append(text)

    table_blocks: list[str] = []
    for index, table in enumerate(doc.tables, start=1):
        rows: list[list[str]] = []
        for row in table.rows:
            values = [_safe_text(cell.text).replace("\r", " ").replace("\n", " ") for cell in row.cells]
            while values and not values[-1]:
                values.pop()
            if any(values):
                rows.append(values)
        if rows:
            table_blocks.append(_rows_to_markdown_table(rows, f"Table {index}"))

    if table_blocks:
        if sections:
            sections.append("## Tables")
        sections.extend(table_blocks)

    return "\n\n".join(sections).strip()


def _extract_doc_text(path: Path) -> str:
    """Best-effort extractor for legacy .doc files."""

    def _try_command(command: list[str]) -> str:
        try:
            completed = subprocess.run(command, capture_output=True, text=True, check=False, timeout=25)
        except Exception:
            return ""
        if completed.returncode != 0:
            return ""
        return completed.stdout.strip()

    command_candidates = (
        ["antiword", str(path)],
        ["catdoc", str(path)],
    )
    for command in command_candidates:
        text = _try_command(command)
        if text:
            return text

    binary = path.read_bytes()
    extracted_chunks: list[str] = []

    try:
        import olefile  # type: ignore

        if olefile.isOleFile(str(path)):
            with olefile.OleFileIO(str(path)) as ole:
                if ole.exists("WordDocument"):
                    binary = ole.openstream("WordDocument").read()
    except Exception:
        pass

    for encoding in ("utf-16le", "utf-8", "gb18030", "latin-1"):
        try:
            decoded = binary.decode(encoding, errors="ignore")
        except Exception:
            continue
        matches = re.findall(r"[A-Za-z0-9\u4e00-\u9fff][A-Za-z0-9\u4e00-\u9fff \t,.;:!?%/()\-_\[\]{}]{8,}", decoded)
        if not matches:
            continue
        extracted_chunks.extend(matches[:2000])
        if len(extracted_chunks) >= 1200:
            break

    normalized_lines: list[str] = []
    seen = set()
    for chunk in extracted_chunks:
        line = re.sub(r"\s+", " ", chunk).strip()
        if len(line) < 6:
            continue
        if line in seen:
            continue
        seen.add(line)
        normalized_lines.append(line)
        if len(normalized_lines) >= 800:
            break

    return "\n".join(normalized_lines).strip()


def _load_spreadsheet_sheets(path: Path) -> list[dict[str, Any]]:
    suffix = path.suffix.lower()
    sheets: list[dict[str, Any]] = []

    if suffix in (".xlsx", ".xlsm"):
        try:
            from openpyxl import load_workbook
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("Missing dependency: openpyxl") from exc

        workbook = load_workbook(filename=str(path), data_only=True, read_only=True)
        try:
            for index, worksheet in enumerate(workbook.worksheets, start=1):
                rows: list[list[str]] = []
                for row in worksheet.iter_rows(values_only=True):
                    values = [_safe_text(cell) for cell in row]
                    while values and not values[-1]:
                        values.pop()
                    if any(values):
                        rows.append(values)
                sheets.append({"page": index, "name": worksheet.title or f"Sheet{index}", "rows": rows})
        finally:
            workbook.close()
        return sheets

    if suffix == ".xls":
        try:
            import xlrd
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("Missing dependency: xlrd") from exc

        workbook = xlrd.open_workbook(str(path), on_demand=True)
        try:
            for index in range(workbook.nsheets):
                worksheet = workbook.sheet_by_index(index)
                rows: list[list[str]] = []
                for row_idx in range(worksheet.nrows):
                    values = [_safe_text(worksheet.cell_value(row_idx, col_idx)) for col_idx in range(worksheet.ncols)]
                    while values and not values[-1]:
                        values.pop()
                    if any(values):
                        rows.append(values)
                sheet_name = workbook.sheet_names()[index] if index < len(workbook.sheet_names()) else f"Sheet{index + 1}"
                sheets.append({"page": index + 1, "name": sheet_name or f"Sheet{index + 1}", "rows": rows})
        finally:
            workbook.release_resources()
        return sheets

    return sheets


def _extract_spreadsheet_pages(path: Path) -> list[dict[str, Any]]:
    sheets = _load_spreadsheet_sheets(path)
    pages: list[dict[str, Any]] = []
    for sheet in sheets:
        rows = sheet["rows"]
        name = str(sheet["name"])
        normalized = _normalize_tabular_rows(rows)
        headers = list(normalized["headers"])
        body = list(normalized["body"])
        header_row_number = int(normalized["header_row_number"])
        data_start_row = int(normalized["data_start_row"])
        row_count = int(normalized["row_count"])
        column_count = int(normalized["column_count"])

        limited_headers = headers[:TABLE_VIEW_MAX_COLUMNS]
        limited_rows = [row[:TABLE_VIEW_MAX_COLUMNS] for row in body[:TABLE_VIEW_MAX_ROWS]]
        table_truncated = len(body) > TABLE_VIEW_MAX_ROWS or len(headers) > TABLE_VIEW_MAX_COLUMNS

        markdown = _rows_to_markdown_table(rows, f"Sheet: {name}")
        plain_text = _rows_to_plain_records(rows, name)
        pages.append(
            {
                "page": int(sheet["page"]),
                "name": name,
                "text_markdown": markdown,
                "text_plain": plain_text,
                "table_headers": limited_headers,
                "table_rows": limited_rows,
                "table_data_start_row": data_start_row,
                "table_header_row_number": header_row_number,
                "table_total_rows": row_count,
                "table_total_columns": column_count,
                "table_truncated": table_truncated,
            }
        )

    if not pages:
        pages.append(
            {
                "page": 1,
                "name": "Sheet1",
                "text_markdown": "### Sheet: Sheet1\n\n_No rows in this sheet._",
                "text_plain": "Sheet: Sheet1\n(empty)",
                "table_headers": [],
                "table_rows": [],
                "table_data_start_row": 1,
                "table_header_row_number": 0,
                "table_total_rows": 0,
                "table_total_columns": 0,
                "table_truncated": False,
            }
        )
    return pages


def _extract_delimited_text_as_markdown(path: Path, delimiter: str) -> str:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    reader = csv.reader(io.StringIO(raw), delimiter=delimiter)
    rows: list[list[str]] = []
    for row in reader:
        values = [_safe_text(cell) for cell in row]
        while values and not values[-1]:
            values.pop()
        if any(values):
            rows.append(values)
    return _rows_to_markdown_table(rows, "Delimited Preview")


def read_file_page_text(path: Path, page: int | None = None) -> dict[str, Any]:
    """Return normalized text payload for the unified frontend viewer."""
    suffix = path.suffix.lower()
    page_number = max(1, page or 1)

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("Missing dependency: pypdf") from exc

        reader = PdfReader(str(path))
        total_pages = len(reader.pages)
        if page_number < 1 or page_number > total_pages:
            raise ValueError(f"page out of range: 1..{total_pages}")

        text = (reader.pages[page_number - 1].extract_text() or "").strip()
        return {
            "page": page_number,
            "page_count": total_pages,
            "page_label": f"Page {page_number}",
            "format": "plain",
            "text": text,
        }

    if suffix in SPREADSHEET_EXTENSIONS:
        pages = _extract_spreadsheet_pages(path)
        total_pages = len(pages)
        if page_number < 1 or page_number > total_pages:
            raise ValueError(f"page out of range: 1..{total_pages}")

        selected = pages[page_number - 1]
        return {
            "page": page_number,
            "page_count": total_pages,
            "page_label": selected["name"],
            "format": "table",
            "text": selected["text_markdown"],
            "table_headers": selected["table_headers"],
            "table_rows": selected["table_rows"],
            "table_data_start_row": selected["table_data_start_row"],
            "table_header_row_number": selected["table_header_row_number"],
            "table_total_rows": selected["table_total_rows"],
            "table_total_columns": selected["table_total_columns"],
            "table_truncated": selected["table_truncated"],
        }

    if suffix == ".docx":
        text = _extract_docx_markdown(path)
        return {
            "page": 1,
            "page_count": 1,
            "page_label": "Document",
            "format": "markdown",
            "text": text,
        }

    if suffix == ".doc":
        text = _extract_doc_text(path)
        if not text:
            text = (
                "Unable to extract rich text from this .doc file in current environment.\n"
                "Please convert to .docx for higher-fidelity preview and retrieval."
            )
        return {
            "page": 1,
            "page_count": 1,
            "page_label": "Document",
            "format": "plain",
            "text": text,
        }

    if suffix == ".csv":
        return {
            "page": 1,
            "page_count": 1,
            "page_label": "Preview",
            "format": "markdown",
            "text": _extract_delimited_text_as_markdown(path, ","),
        }

    if suffix == ".tsv":
        return {
            "page": 1,
            "page_count": 1,
            "page_label": "Preview",
            "format": "markdown",
            "text": _extract_delimited_text_as_markdown(path, "\t"),
        }

    if suffix in TEXT_FILE_EXTENSIONS:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return {
            "page": 1,
            "page_count": 1,
            "page_label": "Preview",
            "format": "markdown" if suffix in MARKDOWN_EXTENSIONS else "plain",
            "text": text,
        }

    raise ValueError(f"Unsupported extension: {suffix}")


def load_documents_from_file(path: Path) -> list[Document]:
    """Convert a single source file into LangChain Documents."""
    suffix = path.suffix.lower()
    base_metadata = {
        "source": _relative_source(path),
        "file_name": path.name,
        "extension": suffix,
    }

    if suffix in TEXT_FILE_EXTENSIONS:
        if suffix == ".csv":
            text = _extract_delimited_text_as_markdown(path, ",")
        elif suffix == ".tsv":
            text = _extract_delimited_text_as_markdown(path, "\t")
        else:
            text = path.read_text(encoding="utf-8", errors="ignore").strip()
        if not text:
            return []
        return [Document(page_content=text, metadata=base_metadata)]

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("Missing dependency: pypdf") from exc

        reader = PdfReader(str(path))
        documents: list[Document] = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not text:
                continue
            metadata = {**base_metadata, "page": page_number}
            documents.append(Document(page_content=text, metadata=metadata))
        return documents

    if suffix == ".docx":
        text = _extract_docx_markdown(path)
        if not text:
            return []
        return [Document(page_content=text, metadata=base_metadata)]

    if suffix == ".doc":
        text = _extract_doc_text(path)
        if not text:
            return []
        return [Document(page_content=text, metadata=base_metadata)]

    if suffix in SPREADSHEET_EXTENSIONS:
        pages = _extract_spreadsheet_pages(path)
        documents: list[Document] = []
        for item in pages:
            text = str(item["text_plain"]).strip()
            if not text:
                continue
            metadata = {
                **base_metadata,
                "page": int(item["page"]),
                "sheet_name": str(item["name"]),
            }
            documents.append(Document(page_content=text, metadata=metadata))
        return documents

    return []
